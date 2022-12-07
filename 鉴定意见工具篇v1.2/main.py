import docx
# 按 Shift+F10 执行或将其替换为您的代码。
# 按 双击 Shift 在所有地方搜索类、文件、工具窗口、操作和设置。
def read_report(report_name,case_num):  #读取鉴定意见书
    doc = docx.Document(report_name)
    count = len(doc.paragraphs)

    case = ''#案件编号
    demand = [] #委托事项
    opinions = [] #鉴定意见
    Evaluators = '' #鉴定人
    Date = '' #鉴定日期
    standard=[] # 原始记录开头规范
    mater_dic = {} #检材编号及检材类型

    for i in range(count):
        if '委托鉴定事项' in doc.paragraphs[i].text:
            demand_start = i+1
        if '鉴定意见' in doc.paragraphs[i].text:
            opinions_start = i+1
        if '网神洞鉴[' in doc.paragraphs[i].text and ']数鉴字第' in doc.paragraphs[i].text:
            tmp_case = doc.paragraphs[i].text
            case = tmp_case.split('：')[1]
        if '鉴定人员：' in doc.paragraphs[i].text and Evaluators == '':
            Evaluators = doc.paragraphs[i+1].text
        if '鉴定日期' in doc.paragraphs[i].text and Date == '':
            Date = doc.paragraphs[i+1].text
        if '采用的技术标准、技术规范或者技术方法' in doc.paragraphs[i].text:
            standard_start = i+1

    while('受理日期' not in doc.paragraphs[demand_start].text):
        demand.append(doc.paragraphs[demand_start].text)
        demand_start = demand_start + 1
    while('附件' not in doc.paragraphs[opinions_start].text):
        if '此页以下为空白' not in doc.paragraphs[opinions_start].text and doc.paragraphs[opinions_start].text != '':
            opinions.append(doc.paragraphs[opinions_start].text)
        opinions_start = opinions_start + 1
    while('鉴定过程' not in doc.paragraphs[standard_start].text):
        if '此页以下为空白' not in doc.paragraphs[opinions_start].text and doc.paragraphs[opinions_start].text != '':
            standard.append(doc.paragraphs[standard_start].text)
        standard_start = standard_start + 1
    table = doc.tables[-2]
    mater_list = []
    mater = []
    name_type = []
    for row in table.rows:
        for cell in row.cells:
            mater_list.append(cell.text)
    for i in range(len(mater_list)):
        if case_num + '-' in mater_list[i] and mater_list[i] not in mater:
            mater.append(mater_list[i])
        if mater_list[i] == '名称':
            name_type.append(mater_list[i+1])
    for i in range(len(mater)):
        mater_dic[mater[i]] = name_type[i]

    return case,demand,opinions,Evaluators,Date,standard,mater_dic

def run():
    case_num = '2022-150' #此处输入案件号
    report_name = '2022-150鉴定意见书.docx' #此处输入鉴定意见文件的位置
    """
    case = ''#案件编号
    demand = [] #委托事项
    opinions = [] #鉴定意见
    Evaluators = '' #鉴定人
    Date = '' #鉴定日期
    standard=[] # 原始记录开头规范
    mater_dic = {} #检材编号及检材类型
    """
    case,demand,opinions,Evaluators,Date,standard,mater_dic = read_report(report_name,case_num)
    print(case)
    print(demand)
    print(opinions)
    print(Evaluators)
    print(Date)
    print(standard)
    print(mater_dic)


# 按间距中的绿色按钮以运行脚本。
if __name__ == '__main__':
    run()

