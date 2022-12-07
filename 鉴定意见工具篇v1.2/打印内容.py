import docx
import re


def getText(filename):
    doc = docx.Document(filename)
    count = len(doc.paragraphs)
    text=[]
    for i in range(count):
        text.append(doc.paragraphs[i].text)
        print(doc.paragraphs[i].text)
    print(text)
def main():
    getText('2022-150鉴定意见书.docx')


if __name__ == "__main__":
    main()