import docx
import re
import tools
import os
import datetime,time
"""
鉴定意见内容初步筛查
"""
def getText(filename):#获取鉴定意见内容，返回字典类型数据
    doc = docx.Document(filename)
    count = len(doc.paragraphs)

    case=''#案件编号，形如编号：网神洞鉴[2022]数鉴字第XXX号
    number_case=''#案件的数字编号，形如2022-xxx
    date_acceptance_=0#受理日期的索引
    date_acceptance=''#受理日期，形如2022年1月1日
    basic_case=''#基本案情
    ident_matter_=0#委托鉴定事项的索引
    ident_matter=[]#委托鉴定事项
    Evaluators=''#鉴定人员
    date_evaluat=''#鉴定日期
    standard=''#鉴定标准和规范
    instrument=''#检验使用的仪器设备
    prepare_begin=''#本次鉴定开始前....
    Identification_process_=0#鉴定过程的索引
    Identification_process=[]#鉴定过程
    identy=[]#数据保存及数字签名
    identy_=0#数据保存及数字签名的索引
    Analysis_Description_=0#分析说明的索引
    Appraisal_opinions_=0#鉴定意见的索引
    Analysis_Description=[]#分析说明
    Appraisal_opinions=[]#鉴定意见
    annex_=0#附件的下标
    Data_carriers=''#检出数据使用的载体（一般为光盘或者U盘）
    final_date=tools.tools.ChineseToDate(doc.paragraphs[count-1].text)

    for i in range(count):
        if '编号：网神洞鉴' in doc.paragraphs[i].text:
            case=doc.paragraphs[i].text
            case_list = re.findall(r'[1-9]+\.?[0-9]*', case)
            number_case=case_list[0]+'-'+case_list[1]
        if '受理日期' in doc.paragraphs[i].text:
            date_acceptance_=i
            date_acceptance=doc.paragraphs[i+1].text
        if '委托鉴定事项' in doc.paragraphs[i].text:
            ident_matter_=i
        if '基本案情' in doc.paragraphs[i].text:
            basic_case=doc.paragraphs[i+1].text
        if '鉴定人员' in doc.paragraphs[i].text:
            Evaluators=doc.paragraphs[i+1].text
        if '鉴定日期' in doc.paragraphs[i].text:
            date_evaluat=doc.paragraphs[i+1].text
        if '采用的技术标准、技术规范或者技术方法' in doc.paragraphs[i].text:
            standard=doc.paragraphs[i+1].text
        if '检验使用的仪器设备' in doc.paragraphs[i].text:
            instrument=doc.paragraphs[i+1].text
            prepare_begin=doc.paragraphs[i+2].text
        if '鉴定过程' in doc.paragraphs[i].text:
            Identification_process_=i
        if '数据保存及数字签名' in doc.paragraphs[i].text:
            identy_=i
            identy.append(doc.paragraphs[i+1].text)
            identy.append(doc.paragraphs[i + 2].text)
        if '分析说明' in doc.paragraphs[i].text:
            Analysis_Description_=i
        if '鉴定意见' in doc.paragraphs[i].text:
            Appraisal_opinions_=i
        if '附件' in doc.paragraphs[i].text:
            annex_=i
        if '检出数据的' in doc.paragraphs[i].text:
            Data_carriers=doc.paragraphs[i].text
    for i in range(ident_matter_+1,date_acceptance_):
        ident_matter.append(doc.paragraphs[i].text)
    for i in range(Identification_process_+1,identy_):
        Identification_process.append(doc.paragraphs[i].text)
    for i in range(Analysis_Description_+1,Appraisal_opinions_):
        Analysis_Description.append(doc.paragraphs[i].text)
    for i in range(Appraisal_opinions_+1,annex_):
        Appraisal_opinions.append(doc.paragraphs[i].text)
    text_dict={}
    text_dict.setdefault('编号',number_case)
    text_dict.setdefault('受理日期',date_acceptance )
    text_dict.setdefault('基本案情',basic_case )
    text_dict.setdefault('委托鉴定事项',ident_matter )
    text_dict.setdefault('鉴定人员',Evaluators )
    text_dict.setdefault('鉴定日期',date_evaluat )
    text_dict.setdefault('采用的技术标准、技术规范或者技术方法', standard)
    text_dict.setdefault('检验使用的仪器设备', instrument)
    text_dict.setdefault('开始前的准备',prepare_begin )
    text_dict.setdefault('保存和签名',identy )
    text_dict.setdefault('检出数据存放',Data_carriers )
    text_dict.setdefault('落款日期',final_date )
    text_dict.setdefault('鉴定过程',Identification_process )
    text_dict.setdefault('分析说明',Analysis_Description)
    text_dict.setdefault('鉴定意见',Appraisal_opinions )
    return text_dict
def write_context(log_,context,name):
    """
    传过来鉴定过程、分析说明、鉴定意见，筛选出空格和一下此页空白，输入到txt文档，方便填写司鉴系统粘贴
    :param log_:
    :param context:
    :param name:
    :return:
    """
    while '' in context:
        context.remove('')
    for line in context:
        if tools.remove_blank(line):
            context.remove(line)
    with open(log_+'/'+name+'.txt','a+',encoding='utf-8') as write:
        for text in context:
            write.writelines(text+'\n')
def main():
    #修改案件编号和鉴定意见的位置
    case='2022-362'
    text_dict=getText('../鉴定意见案例/2022-362鉴定意见书V2wd.docx')
    now_date,time_name=tools.get_log_name()
    path='./'+case+'鉴定过程&分析说明&鉴定意见'
    log_=path+ time_name
    if not os.path.exists(log_):
        os.makedirs(log_)
    write_context(log_,text_dict.get('鉴定过程'),'鉴定过程')
    write_context(log_,text_dict.get('分析说明'),'分析说明')
    write_context(log_,text_dict.get('鉴定意见'),'鉴定意见')
if     __name__=="__main__":
    main()